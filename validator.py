import os
import re
import json
import argparse
import docx
from docx.shared import Pt

KNOWN_PASSAGES = ["grandmother's loom", "sailing by the stars", "the rooftop sky", "the lot on maple street"]

def count_syllables(word):
    """Establishes syllable count for a word (standard English heuristic)."""
    word = word.lower().strip(".,!?;:()[]\"'")
    if not word:
        return 0
    
    # Common contractions/short words
    if word in ["the", "a", "an", "and", "of", "to"]:
        return 1
        
    vowels = "aeiouy"
    count = 0
    
    # Check first letter
    if word[0] in vowels:
        count += 1
        
    # Check vowel transitions
    for index in range(1, len(word)):
        if word[index] in vowels and word[index - 1] not in vowels:
            count += 1
            
    # Adjustments
    if word.endswith("e"):
        count -= 1
    if word.endswith("le") and len(word) > 2 and word[-3] not in vowels:
        count += 1
        
    # Guard against 0 syllables
    return max(1, count)

def compute_readability(text):
    """Computes Flesch Reading Ease and Flesch-Kincaid Grade Level."""
    # Clean and split sentences
    sentences = re.split(r'[.!?]+(?=\s|$)', text)
    sentences = [s.strip() for s in sentences if s.strip()]
    
    # Split words
    words = re.findall(r'\b[a-zA-Z]+-?[a-zA-Z]*\b', text)
    
    num_sentences = len(sentences)
    num_words = len(words)
    
    if num_sentences == 0 or num_words == 0:
        return {"ease": 0.0, "grade": 0.0, "words": 0, "sentences": 0}
        
    num_syllables = sum(count_syllables(w) for w in words)
    
    ease = 206.835 - 1.015 * (num_words / num_sentences) - 84.6 * (num_syllables / num_words)
    grade = 0.39 * (num_words / num_sentences) + 11.8 * (num_syllables / num_words) - 15.59
    
    return {
        "ease": round(max(0.0, min(100.0, ease)), 1),
        "grade": round(max(0.0, grade), 1),
        "words": num_words,
        "sentences": num_sentences
    }

def detect_placeholders(text):
    """Finds placeholders like [Subject], ##, etc. in a text string."""
    results = []
    for match in re.finditer(r'\[([^\]]+)\]', text):
        results.append(match.group(0))
    for match in re.finditer(r'##', text):
        results.append("##")
    return results

def count_sentences(text):
    """Counts sentences in text using a simple regex tokenizer."""
    clean_text = re.sub(r'\b(e\.g\.|i\.e\.|vs\.|Mr\.|Mrs\.|Dr\.|St\.|Col\.|Row\.|[A-Z]\.)', '', text)
    sentences = re.split(r'[.!?]+(?:\s+|$)(?=[A-Z"]|$)', clean_text)
    sentences = [s.strip() for s in sentences if len(s.strip()) > 3]
    return len(sentences), sentences

def extract_length_constraints(prompt_text):
    """Extracts required sentence length from prompt text (e.g. 'at least 8 complete sentences')."""
    m_at_least = re.search(r'at least\s+(\d+)\s+(?:complete\s+)?sentences', prompt_text, re.IGNORECASE)
    if m_at_least:
        return {"min": int(m_at_least.group(1)), "max": None, "desc": f"At least {m_at_least.group(1)} sentences"}
        
    m_range = re.search(r'(\d+)\s*(?:-|to)\s*(\d+)\s+(?:complete\s+)?sentences', prompt_text, re.IGNORECASE)
    if m_range:
        return {"min": int(m_range.group(1)), "max": int(m_range.group(2)), "desc": f"{m_range.group(1)} to {m_range.group(2)} sentences"}
        
    return None

def convert_to_smart_quotes(text):
    """Replaces straight quotes with smart curly quotes."""
    # Convert double quotes
    text = re.sub(r'^"', '“', text)
    text = re.sub(r'(?<=\s)"', '“', text)
    text = re.sub(r'(?<=[({\[-])"', '“', text)
    text = text.replace('"', '”')
    
    # Convert single quotes
    text = re.sub(r"^'", '‘', text)
    text = re.sub(r"(?<=\s)'", '‘', text)
    text = re.sub(r"(?<=[({\[-])'", '‘', text)
    text = text.replace("'", '’')
    return text

def auto_correct_document(input_path, output_path):
    """Automates font name/size corrections, styles mapping, double spaces, and smart quotes."""
    doc = docx.Document(input_path)
    
    # 1. Delete internal metadata, vendor notes and delivery instructions
    for idx in range(len(doc.paragraphs) - 1, -1, -1):
        p = doc.paragraphs[idx]
        p_low = p.text.lower().strip()
        if ("vendor created" in p_low or 
            "audio transcript from" in p_low or 
            "audio delivery:" in p_low):
            p_element = p._element
            p_element.getparent().remove(p_element)
            p._p = p._element = None
            
    # 2. Process remaining Paragraphs
    for idx, p in enumerate(doc.paragraphs):
        text = p.text
        if not text.strip():
            continue
            
        p_low = text.lower().strip()
        
        # Save original run sizes first before style overrides them
        original_sizes = []
        for run in p.runs:
            original_sizes.append(run.font.size)
        
        # Detect paragraph heading types
        is_main_title = (idx < 5 and ("unit" in p_low or "test" in p_low or "graded assignment" in p_low or "author study" in p_low or idx == 0))
        is_subheading = (
            p_low.startswith("passage ") or 
            p_low.startswith("audio transcript") or 
            p_low == "read the passage." or
            p_low == "read the passages. then, answer the questions." or
            p_low == "refer to the audio and audio script. then, answer the question." or
            p_low.replace('’', "'").replace('‘', "'") in KNOWN_PASSAGES
        )
        
        # Apply standard styles
        try:
            if is_main_title:
                p.style = doc.styles['Heading 2']
            elif is_subheading:
                p.style = doc.styles['Heading 3']
            else:
                if not p.style.name.startswith('Heading'):
                    p.style = doc.styles['Normal']
        except:
            # Fallback if style name doesn't exist in file
            pass
            
        # Clean runs and restore font sizes
        for run_idx, run in enumerate(p.runs):
            run.text = run.text.replace("  ", " ")
            run.text = convert_to_smart_quotes(run.text)
            # Standardized wording replacements
            run.text = run.text.replace("Refer to the audio", "Listen to the audio")
            run.text = run.text.replace("refer to the audio", "listen to the audio")
            run.text = run.text.replace("Read the paragraph", "Read the passage")
            run.text = run.text.replace("read the paragraph", "read the passage")
            
            # Apply Arial standard font
            run.font.name = 'Arial'
            
            # Restore original font size if it was set
            if run_idx < len(original_sizes) and original_sizes[run_idx] is not None:
                run.font.size = original_sizes[run_idx]

    # Process Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.text = run.text.replace("  ", " ")
                        run.text = convert_to_smart_quotes(run.text)
                        # Standardized wording replacements
                        run.text = run.text.replace("Refer to the audio", "Listen to the audio")
                        run.text = run.text.replace("refer to the audio", "listen to the audio")
                        run.text = run.text.replace("Read the paragraph", "Read the passage")
                        run.text = run.text.replace("read the paragraph", "read the passage")
                        run.font.name = 'Arial'
                        
    doc.save(output_path)

class DocumentValidator:
    def __init__(self, filepath):
        self.filepath = filepath
        self.filename = os.path.basename(filepath)
        self.doc = docx.Document(filepath)
        self.doc_type = self._detect_doc_type()
        self.KNOWN_PASSAGES = ["grandmother's loom", "sailing by the stars", "the rooftop sky", "the lot on maple street"]
        
        self.findings = []
        self.stats = {
            "paragraphs_count": len(self.doc.paragraphs),
            "tables_count": len(self.doc.tables),
            "placeholders_count": 0,
            "font_violations_count": 0,
            "wording_violations_count": 0,
            "grammar_violations_count": 0,
            "overall_compliance_score": 100
        }
        self.html_preview = []  # List of elements for in-context rendering

    def _detect_doc_type(self):
        if "answer" in self.filename.lower() or "key" in self.filename.lower():
            return "part2_answer_key"
        
        full_text = " ".join([p.text for p in self.doc.paragraphs])
        
        if "part 1" in self.filename.lower():
            return "part1_test"
        if "part 2" in self.filename.lower():
            return "part2_test"
            
        if "part 1" in full_text.lower():
            return "part1_test"
        if "part 2" in full_text.lower():
            return "part2_test"
            
        return "part1_test"

    def add_finding(self, category, level, message, element_type, index, text_snippet=None):
        finding = {
            "category": category,      # "font", "placeholder", "wording", "grammar"
            "level": level,            # "error", "warning", "info"
            "message": message,
            "element_type": element_type,
            "index": index,
            "text": text_snippet
        }
        self.findings.append(finding)
        
        if category == "font":
            self.stats["font_violations_count"] += 1
        elif category == "placeholder":
            self.stats["placeholders_count"] += 1
        elif category == "wording":
            self.stats["wording_violations_count"] += 1
        elif category == "grammar":
            self.stats["grammar_violations_count"] += 1

    def run_validation(self):
        # 1. Paragraph Checks (Fonts, placeholders, double spaces)
        for idx, p in enumerate(self.doc.paragraphs):
            text = p.text
            style_name = p.style.name
            
            if not text.strip():
                self.html_preview.append({"type": "paragraph", "style": style_name, "text": "", "findings": []})
                continue
                
            p_findings = []
            
            # Placeholders Check
            placeholders = detect_placeholders(text)
            for ph in placeholders:
                msg = f"Placeholder detected: '{ph}'"
                self.add_finding("placeholder", "error", msg, "paragraph", idx, text[:120])
                p_findings.append({"category": "placeholder", "level": "error", "message": msg})
                
            p_low = text.lower().strip()
            
            # Check for vendor notes or audio delivery metadata
            if ("vendor created" in p_low or 
                "audio transcript from" in p_low or 
                "audio delivery:" in p_low):
                msg = f"Internal metadata or vendor note detected: '{text}'. This should not be present in the final document."
                self.add_finding("wording", "error", msg, "paragraph", idx, text[:120])
                p_findings.append({"category": "wording", "level": "error", "message": msg})
                
            # Font Check
            
            # Check Title
            if idx == 0 or (idx < 5 and ("unit" in p_low or "test" in p_low or "graded assignment" in p_low or "author study" in p_low)):
                expected_styles = ["Heading 1", "Heading 2", "Title"]
                if style_name not in expected_styles:
                    msg = f"Title paragraph uses style '{style_name}' instead of standard Heading style (Heading 1 or Heading 2)."
                    self.add_finding("font", "warning", msg, "paragraph", idx, text[:100])
                    p_findings.append({"category": "font", "level": "warning", "message": msg})
                    
            # Check subheadings
            is_sub = (
                p_low.startswith("passage ") or 
                p_low.startswith("audio transcript") or 
                p_low == "read the passage." or
                p_low == "read the passages. then, answer the questions." or
                p_low == "refer to the audio and audio script. then, answer the question." or
                p_low.replace('’', "'").replace('‘', "'") in self.KNOWN_PASSAGES
            )
            if is_sub:
                if style_name != "Heading 3" and style_name != "Heading 2":
                    msg = f"Subheader uses style '{style_name}' instead of standard subheader style (Heading 3)."
                    self.add_finding("font", "warning", msg, "paragraph", idx, text[:100])
                    p_findings.append({"category": "font", "level": "warning", "message": msg})

            # Check individual runs for fonts & sizes
            for run in p.runs:
                f_name = run.font.name
                f_size = run.font.size.pt if run.font.size else None
                
                if f_name and f_name != "Arial":
                    msg = f"Non-conforming font family '{f_name}' detected. Standard font is Arial."
                    self.add_finding("font", "error", msg, "paragraph", idx, run.text[:100])
                    p_findings.append({"category": "font", "level": "error", "message": msg})
                    
                if f_size:
                    is_title = (idx < 5 and ("unit" in p_low or "test" in p_low or "graded assignment" in p_low or "author study" in p_low or idx == 0))
                    if is_title:
                        if f_size < 14.0:
                            msg = f"Title font size too small: {f_size}pt. Expected 16pt-24pt."
                            self.add_finding("font", "info", msg, "paragraph", idx, run.text[:100])
                            p_findings.append({"category": "font", "level": "info", "message": msg})
                    else:
                        if is_sub:
                            if f_size != 11.0:
                                msg = f"Non-conforming subheading font size: {f_size}pt. Expected 11.0pt."
                                self.add_finding("font", "info", msg, "paragraph", idx, run.text[:100])
                                p_findings.append({"category": "font", "level": "info", "message": msg})
                        else:
                            if f_size != 10.0 and f_size != 11.0:
                                msg = f"Non-conforming font size: {f_size}pt. Expected Arial 10.0pt or 11.0pt."
                                self.add_finding("font", "error", msg, "paragraph", idx, run.text[:100])
                                p_findings.append({"category": "font", "level": "error", "message": msg})
                            
            # Grammar, Punctuation & Style Checks
            # Double space check
            if "  " in text:
                msg = "Double space detected."
                self.add_finding("grammar", "warning", msg, "paragraph", idx, text[:120])
                p_findings.append({"category": "grammar", "level": "warning", "message": msg})
                
            # Wording guidelines check
            if "refer to the audio" in p_low:
                msg = "Wording violation: 'refer to the audio' detected. ELA guidelines require 'listen to the audio'."
                self.add_finding("wording", "warning", msg, "paragraph", idx, text[:120])
                p_findings.append({"category": "wording", "level": "warning", "message": msg})
                
            if "read the paragraph" in p_low:
                msg = "Wording violation: 'read the paragraph' detected. ELA guidelines require 'read the passage' or 'read the passages'."
                self.add_finding("wording", "warning", msg, "paragraph", idx, text[:120])
                p_findings.append({"category": "wording", "level": "warning", "message": msg})
                
            # Straight quote check
            if '"' in text or "'" in text:
                msg = "Straight quotes detected. Standard ELA typography requires smart curly quotes (“ ” or ‘ ’)."
                self.add_finding("grammar", "info", msg, "paragraph", idx, text[:120])
                p_findings.append({"category": "grammar", "level": "info", "message": msg})
                
            # Check for ending punctuation in questions / directions (exclude headers/names)
            is_header = (style_name.startswith("Heading") or "rubric" in p_low or "answer:" in p_low or p_low.startswith("item") or "graded assignment" in p_low or p_low.startswith("unit test") or p_low.startswith("total score") or p_low.startswith("name:") or p_low.startswith("date:"))
            if text.strip() and not is_header:
                last_char = text.strip()[-1]
                if last_char not in [".", "?", "!", '"', '”', '’', ':', '_', ')', ']']:
                    msg = f"Paragraph does not end with terminal punctuation (ends with '{last_char}')."
                    self.add_finding("grammar", "warning", msg, "paragraph", idx, text[-50:])
                    p_findings.append({"category": "grammar", "level": "warning", "message": msg})
                    
            # Check capitalize first letter
            clean_text = text.strip().lstrip("\"'“‘_-#0123456789. ")
            if clean_text and clean_text[0].islower():
                msg = "Sentence starts with a lowercase letter."
                self.add_finding("grammar", "warning", msg, "paragraph", idx, text[:50])
                p_findings.append({"category": "grammar", "level": "warning", "message": msg})

            self.html_preview.append({
                "type": "paragraph",
                "style": style_name,
                "text": text,
                "findings": p_findings
            })

        # 2. Table Checks (Structure, cells, and rubrics)
        for t_idx, table in enumerate(self.doc.tables):
            table_findings = []
            num_rows = len(table.rows)
            num_cols = len(table.columns)
            
            if self.doc_type == "part1_test":
                if num_cols == 3:
                    msg = f"Table {t_idx} has 3 columns instead of 2. In standard templates, metadata tables have 2 columns (Metadata Field, Value)."
                    self.add_finding("wording", "info", msg, "table", t_idx)
                    table_findings.append({"category": "wording", "level": "info", "message": msg})
                elif num_cols != 2 and num_cols != 4:
                    msg = f"Non-conforming table structure: {num_cols} columns. Expected 2 columns for item metadata, or 4 columns for blueprint."
                    self.add_finding("wording", "error", msg, "table", t_idx)
                    table_findings.append({"category": "wording", "level": "error", "message": msg})

            cell_data = []
            
            for r_idx, row in enumerate(table.rows):
                row_cells = []
                for c_idx, cell in enumerate(row.cells):
                    cell_text = cell.text
                    cell_findings = []
                    
                    placeholders = detect_placeholders(cell_text)
                    for ph in placeholders:
                        msg = f"Placeholder detected in table cell: '{ph}'"
                        self.add_finding("placeholder", "error", msg, "table", t_idx, f"Row {r_idx} Col {c_idx}: {cell_text[:80]}")
                        cell_findings.append({"category": "placeholder", "level": "error", "message": msg})
                        
                    for cp in cell.paragraphs:
                        for run in cp.runs:
                            f_name = run.font.name
                            f_size = run.font.size.pt if run.font.size else None
                            
                            if f_name and f_name != "Arial":
                                msg = f"Non-conforming font family '{f_name}' in table cell. Standard is Arial."
                                self.add_finding("font", "error", msg, "table", t_idx, run.text[:100])
                                cell_findings.append({"category": "font", "level": "error", "message": msg})
                                
                            if f_size and f_size != 10.0 and f_size != 11.0:
                                msg = f"Non-conforming font size: {f_size}pt in table. Expected Arial 10.0pt or 11.0pt."
                                self.add_finding("font", "error", msg, "table", t_idx, run.text[:100])
                                cell_findings.append({"category": "font", "level": "error", "message": msg})

                    if "  " in cell_text:
                        msg = "Double space detected in cell."
                        self.add_finding("grammar", "warning", msg, "table", t_idx, cell_text[:100])
                        cell_findings.append({"category": "grammar", "level": "warning", "message": msg})
                        
                    row_cells.append({
                        "text": cell_text,
                        "findings": cell_findings
                    })
                cell_data.append(row_cells)
                
            if self.doc_type == "part2_answer_key" and num_cols == 2 and num_rows == 2:
                r0_col0 = cell_data[0][0]["text"].lower().strip()
                r1_col0 = cell_data[1][0]["text"].lower().strip()
                
                if "rubric" in r0_col0 and "exemplar" in r1_col0:
                    rubric_desc = cell_data[0][1]["text"]
                    exemplar_ans = cell_data[1][1]["text"]
                    
                    question_prompt = ""
                    q_num_str = f"{t_idx + 1}."
                    
                    for p in self.doc.paragraphs:
                        p_txt = p.text.strip()
                        if p_txt.startswith(q_num_str) or p_txt.startswith(f"Item {t_idx + 1}:") or (f"Question {t_idx + 1}" in p_txt):
                            question_prompt = p_txt
                            break
                            
                    if question_prompt:
                        constraints = extract_length_constraints(question_prompt)
                        if constraints:
                            s_count, sentences = count_sentences(exemplar_ans)
                            is_invalid = False
                            if constraints["min"] and s_count < constraints["min"]:
                                is_invalid = True
                            if constraints["max"] and s_count > constraints["max"]:
                                is_invalid = True
                                
                            if is_invalid:
                                msg = f"Exemplar Answer has {s_count} sentences. Prompt requirement: '{constraints['desc']}'."
                                self.add_finding("wording", "error", msg, "table", t_idx, exemplar_ans[:150])
                                table_findings.append({"category": "wording", "level": "error", "message": msg})
                            else:
                                msg = f"Exemplar Answer has {s_count} sentences. Prompt requirement: '{constraints['desc']}' - PASS."
                                self.add_finding("wording", "info", msg, "table", t_idx, exemplar_ans[:150])
                                table_findings.append({"category": "wording", "level": "info", "message": msg})

            self.html_preview.append({
                "type": "table",
                "rows": num_rows,
                "cols": num_cols,
                "cells": cell_data,
                "findings": table_findings
            })

        # 3. Overall Readability Calculations
        full_doc_text = "\n".join([p.text for p in self.doc.paragraphs])
        
        passages = {
            "Grandmother's Loom": "",
            "Sailing by the Stars": "",
            "The Rooftop Sky": "",
            "The Lot on Maple Street": ""
        }
        
        current_passage = None
        for p in self.doc.paragraphs:
            txt = p.text.strip()
            if txt in passages:
                current_passage = txt
            elif current_passage and txt:
                if txt.startswith("Item") or txt.startswith("Unit") or "rubric" in txt.lower() or "answer" in txt.lower():
                    current_passage = None
                else:
                    passages[current_passage] += "\n" + txt
                    
        self.passages_readability = {}
        for p_name, p_text in passages.items():
            if len(p_text.strip()) > 100:
                res = compute_readability(p_text)
                self.passages_readability[p_name] = res
                
                grade_lvl = res["grade"]
                if grade_lvl < 4.0 or grade_lvl > 7.0:
                    msg = f"Passage '{p_name}' Flesch-Kincaid Grade Level is {grade_lvl}. Grade 5 expected range: 4.5 to 6.5."
                    self.add_finding("wording", "info", msg, "section", 0)
                else:
                    msg = f"Passage '{p_name}' Readability: Grade {grade_lvl} ({res['ease']} Ease). Pass."
                    self.add_finding("wording", "info", msg, "section", 0)

        total_deduction = 0
        for f in self.findings:
            if f["level"] == "error":
                total_deduction += 5
            elif f["level"] == "warning":
                total_deduction += 2
        self.stats["overall_compliance_score"] = max(0, 100 - total_deduction)

        return {
            "doc_type": self.doc_type,
            "stats": self.stats,
            "findings": self.findings,
            "passages_readability": self.passages_readability,
            "html_preview": self.html_preview
        }

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Validate ELA document wording and fonts")
    parser.add_argument("--file", type=str, help="Path to docx file")
    parser.add_argument("--test", action="store_true", help="Run test on standard files")
    parser.add_argument("--correct", nargs=2, metavar=('IN', 'OUT'), help="Correct file IN and save to OUT")
    args = parser.parse_args()
    
    if args.correct:
        print(f"Correcting: {args.correct[0]} -> Saving to: {args.correct[1]}")
        auto_correct_document(args.correct[0], args.correct[1])
        print("Corrections completed!")
    elif args.test:
        test_dir = r"C:\Users\Riyak\.gemini\antigravity\scratch\english-paper-validator"
        test_file = os.path.join(test_dir, "ELA_Grade_5_Author_Study_Part_2_Answer_Key.docx")
        if os.path.exists(test_file):
            print(f"Testing validator on: {test_file}")
            validator = DocumentValidator(test_file)
            res = validator.run_validation()
            print("--- STATS ---")
            print(json.dumps(res["stats"], indent=2))
            print("--- READABILITY ---")
            print(json.dumps(res["passages_readability"], indent=2))
        else:
            print(f"Test file not found at {test_file}")
    elif args.file:
        if os.path.exists(args.file):
            validator = DocumentValidator(args.file)
            res = validator.run_validation()
            print(json.dumps(res, indent=2))
        else:
            print(json.dumps({"error": f"File '{args.file}' not found."}))
    else:
        print("Please provide --file, --test, or --correct <IN> <OUT>")
